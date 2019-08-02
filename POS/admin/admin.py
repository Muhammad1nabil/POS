import sqlite3
from collections import OrderedDict
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import os
from kivy.app import App
from kivy.clock import Clock
from kivy.uix.modalview import ModalView
from kivy.uix.spinner import Spinner
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.label import Label
from kivy.uix.button import Button
from kivy.uix.textinput import TextInput
from kivy.lang import Builder
from utils.datatables import DataTable
from utils.translate import Ar_text, arabic_trans
from datetime import datetime
from docx.shared import Cm
Builder.load_file('admin/admin.kv')


class Notify(ModalView):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.size_hint = (0.7, 0.7)


class AdminWindow(BoxLayout):

    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.conn = sqlite3.connect('POS.db')
        self.cur = self.conn.cursor()

        self.notify = Notify()

        content = self.ids.scrn_contents
        users = self.get_users()
        userstable = DataTable(table=users)
        content.add_widget(userstable)

        # Display Products
        product_scrn = self.ids.scrn_product_contents
        products = self.get_products()
        self.prodkeys = []
        self.prodnames = []
        self.prod_ids = list(products['product_code'].values())
        for prod_key, value in products['product_name'].items():
            self.prodnames.append(value)
            self.prodkeys.append(prod_key)

        prod_table = DataTable(table=products)
        product_scrn.add_widget(prod_table)

        # Display Customers
        customer_scrn = self.ids.scrn_customer_contents
        customers = self.get_customers()
        cust_table = DataTable(table=customers)
        customer_scrn.add_widget(cust_table)

        # Display Orders
        orders_scrn = self.ids.scrn_orders_contents
        orders = self.get_orders()
        order_table = DataTable(table=orders)
        orders_scrn.add_widget(order_table)

        # stats data
        self.current_order_id_details = ''

    def logout(self):
        self.parent.parent.current = 'scrn_si'

    def killswitch(self, dtx):
        self.notify.dismiss()
        self.notify.clear_widgets()

    def add_user_fields(self):
        target = self.ids.ops_fields
        target.clear_widgets()
        crud_first = TextInput(hint_text='First Name', multiline=False)
        crud_last = TextInput(hint_text='Last Name', multiline=False)
        crud_user = TextInput(hint_text='User Name', multiline=False)
        crud_pwd = TextInput(hint_text='Password', multiline=False)
        crud_des = Spinner(text='Operator', values=['Operator', 'Administrator'])
        crud_submit = Button(text='Add', size_hint_x=None, width=100,
                             on_release=lambda x: self.add_user(crud_first.text, crud_last.text, crud_user.text,
                                                                crud_pwd.text, crud_des.text))

        target.add_widget(crud_first)
        target.add_widget(crud_last)
        target.add_widget(crud_user)
        target.add_widget(crud_pwd)
        target.add_widget(crud_des)
        target.add_widget(crud_submit)

    def update_user_fields(self):
        target = self.ids.ops_fields
        target.clear_widgets()
        crud_first = TextInput(hint_text='First Name', multiline=False)
        crud_last = TextInput(hint_text='Last Name', multiline=False)
        crud_user = TextInput(hint_text='User Name', multiline=False)
        crud_pwd = TextInput(hint_text='Password', multiline=False)
        crud_des = Spinner(text='Operator', values=['Operator', 'Administrator'])
        crud_submit = Button(text='Update', size_hint_x=None, width=100,
                             on_release=lambda x: self.update_user(crud_first.text, crud_last.text, crud_user.text,
                                                                   crud_pwd.text, crud_des.text))

        target.add_widget(crud_first)
        target.add_widget(crud_last)
        target.add_widget(crud_user)
        target.add_widget(crud_pwd)
        target.add_widget(crud_des)
        target.add_widget(crud_submit)

    def remove_user_fields(self):
        target = self.ids.ops_fields
        target.clear_widgets()
        crud_user = TextInput(hint_text='User Name', multiline=False)
        crud_submit = Button(text='Remove', size_hint_x=None, width=100,
                             on_release=lambda x: self.remove_user(crud_user.text))

        target.add_widget(crud_user)
        target.add_widget(crud_submit)

    def add_user(self, first, last, user, pwd, des):
        content = self.ids.scrn_contents
        content.clear_widgets()

        if first == '' or last == '' or user == '' or pwd == '':
            self.notify.add_widget(Label(text='[color=#FF0000][b]All Fields Required![/b][/color]', markup=True))
            self.notify.open()
            Clock.schedule_once(self.killswitch, 1)
        else:
            query = 'INSERT INTO USERS(First_Name, Last_Name, Username, Password, Designation) VALUES(?, ?, ?, ?, ?)'
            values = [first, last, user, pwd, des]
            self.cur.execute(query, values)
            self.conn.commit()

        users = self.get_users()
        userstable = DataTable(table=users)
        content.add_widget(userstable)
        self.add_user_fields()

    def update_user(self, first, last, user, pwd, des):
        content = self.ids.scrn_contents
        content.clear_widgets()

        if user == '':
            self.notify.add_widget(Label(text='[color=#FF0000][b]Username is Required![/b][/color]', markup=True))
            self.notify.open()
            Clock.schedule_once(self.killswitch, 1)
        else:
            users = dict(self.get_users().items())
            user_names = users.get('user_names')
            usernamelist = list(dict(user_names).values())
            x = usernamelist.__contains__(user)
            if not x:
                self.notify.add_widget(Label(text='[color=#FF0000][b]Invalid Username![/b][/color]', markup=True))
                self.notify.open()
                Clock.schedule_once(self.killswitch, 1)
            else:
                index = usernamelist.index(user)
                if first == '':
                    first_names = users.get('first_names')
                    firstnamelist = list(dict(first_names).values())
                    first = firstnamelist[index]
                if last == '':
                    last_names = users.get('last_names')
                    lastnamelist = list(dict(last_names).values())
                    last = lastnamelist[index]
                if pwd == '':
                    pwds = users.get('passwords')
                    pwdlist = list(dict(pwds).values())
                    pwd = pwdlist[index]

                query = 'UPDATE USERS SET First_Name=?, Last_Name=?, Username=?, Password=?, Designation=? WHERE Username=?'
                values = [first, last, user, pwd, des, user]
                self.cur.execute(query, values)
                self.conn.commit()

        users = self.get_users()
        userstable = DataTable(table=users)
        content.add_widget(userstable)
        self.update_user_fields()

    def remove_user(self, user):
        content = self.ids.scrn_contents
        content.clear_widgets()

        if user == '':
            self.notify.add_widget(Label(text='[color=#FF0000][b]Username is Required![/b][/color]', markup=True))
            self.notify.open()
            Clock.schedule_once(self.killswitch, 1)
        else:
            users = dict(self.get_users().items())
            user_names = users.get('user_names')
            usernamelist = list(dict(user_names).values())
            x = usernamelist.__contains__(user)
            if not x:
                self.notify.add_widget(Label(text='[color=#FF0000][b]Invalid Username![/b][/color]', markup=True))
                self.notify.open()
                Clock.schedule_once(self.killswitch, 1)
            else:
                query = 'DELETE FROM USERS WHERE Username=?'
                values = [user]
                self.cur.execute(query, values)
                self.conn.commit()

        users = self.get_users()
        userstable = DataTable(table=users)
        content.add_widget(userstable)
        self.remove_user_fields()

    def get_users(self):
        _users = OrderedDict()
        _users['first_names'] = {}
        _users['last_names'] = {}
        _users['user_names'] = {}
        _users['passwords'] = {}
        _users['designations'] = {}
        first_names = []
        last_names = []
        user_names = []
        passwords = []
        designations = []

        query = 'SELECT * FROM USERS'
        self.cur.execute(query)
        users = self.cur.fetchall()

        for user in users:
            first_names.append(user[1])
            last_names.append(user[2])
            user_names.append(user[3])
            pwd = user[4]
            if len(pwd) > 10:
                pwd = pwd[:10] + '...'
            passwords.append(pwd)
            designations.append(user[5])
        users_length = len(first_names)
        idx = 0
        while idx < users_length:
            _users['first_names'][idx] = first_names[idx]
            _users['last_names'][idx] = last_names[idx]
            _users['user_names'][idx] = user_names[idx]
            _users['passwords'][idx] = passwords[idx]
            _users['designations'][idx] = designations[idx]

            idx += 1

        return _users

    # ----------------------------------------------------------------------------------
    # ----------------------------------------------------------------------------------
    # ----------------------------------------------------------------------------------
    # ----------------------------------------------------------------------------------

    def add_product_fields(self):
        target = self.ids.ops_fields_p
        target.clear_widgets()

        crud_code = TextInput(hint_text='Product Code', multiline=False)
        crud_name = Ar_text(hint_text='Product Name', multiline=False, font_name="utils/STC")
        crud_price = TextInput(hint_text='Price', multiline=False, input_filter="float")
        crud_cat = Spinner(text=arabic_trans('النوع'), font_name="utils/STC", values=[arabic_trans('خضار'), arabic_trans('فاكهة'), arabic_trans('منتجات'),
                                                               arabic_trans('مأكولات جاهزه')])
        crud_submit = Button(text='Add', size_hint_x=None, width=100,
                             on_release=lambda x: self.add_product(crud_code.text, crud_name.text, crud_price.text,
                                                                   crud_cat.text))

        target.add_widget(crud_code)
        target.add_widget(crud_name)
        target.add_widget(crud_price)
        target.add_widget(crud_cat)
        target.add_widget(crud_submit)

    def update_product_fields(self):
        target = self.ids.ops_fields_p
        target.clear_widgets()

        crud_code = TextInput(hint_text='Product Code', multiline=False, font_name="utils/STC")
        crud_name = Ar_text(hint_text='Product Name', multiline=False, font_name="utils/STC")
        crud_price = TextInput(hint_text='Price', multiline=False, input_filter="float")
        crud_cat = Spinner(text=arabic_trans('النوع'), font_name="utils/STC", values=[arabic_trans('خضار'), arabic_trans('فاكهة'), arabic_trans('منتجات'),
                                                               arabic_trans('مأكولات جاهزه')])
        crud_submit = Button(text='Update', size_hint_x=None, width=100,
                             on_release=lambda x: self.update_product(crud_code.text, crud_name.text, crud_price.text,
                                                                      crud_cat.text))

        target.add_widget(crud_code)
        target.add_widget(crud_name)
        target.add_widget(crud_price)
        target.add_widget(crud_cat)
        target.add_widget(crud_submit)

    def remove_product_fields(self):
        target = self.ids.ops_fields_p
        target.clear_widgets()
        crud_code = TextInput(hint_text='Product Code', multiline=False)
        crud_submit = Button(text='Remove', size_hint_x=None, width=100,
                             on_release=lambda x: self.remove_product(crud_code.text))

        target.add_widget(crud_code)
        target.add_widget(crud_submit)

    def add_product(self, code, name, price, category):
        content = self.ids.scrn_product_contents
        content.clear_widgets()

        if code == '' or name == '' or price == '' or category == arabic_trans('النوع'):
            self.notify.add_widget(Label(text='[color=#FF0000][b]All Fields Required![/b][/color]', markup=True))
            self.notify.open()
            Clock.schedule_once(self.killswitch, 1)
        else:
            query = 'INSERT INTO PRODUCTS(product_code, product_name, price, Category) VALUES(?, ?, ?, ?)'
            values = [str(code), arabic_trans(name), price, arabic_trans(category)]
            self.cur.execute(query, values)
            self.conn.commit()

        prodz = self.get_products()
        stocktable = DataTable(table=prodz)
        content.add_widget(stocktable)
        self.add_product_fields()

    def update_product(self, code, name, price, cat):
        content = self.ids.scrn_product_contents
        content.clear_widgets()

        if code == '':
            self.notify.add_widget(Label(text='[color=#FF0000][b]Product code is Required![/b][/color]', markup=True))
            self.notify.open()
            Clock.schedule_once(self.killswitch, 1)
        else:
            products = dict(self.get_products().items())
            cod = products.get('product_code')
            codelist = list(str(a) for a in dict(cod).values())
            x = codelist.__contains__(code)
            if not x:
                self.notify.add_widget(Label(text='[color=#FF0000][b]Invalid Product Code![/b][/color]', markup=True))
                self.notify.open()
                Clock.schedule_once(self.killswitch, 1)
            else:
                index = codelist.index(code)
                if name == '':
                    names = products.get('product_name')
                    namelist = list(dict(names).values())
                    name = namelist[index]
                else:
                    name = arabic_trans(name)
                if price == '':
                    prices = products.get('price')
                    pricelist = list(dict(prices).values())
                    price = pricelist[index]
                if cat == arabic_trans('النوع'):
                    cats = products.get('Category')
                    catlist = list(dict(cats).values())
                    cat = arabic_trans(catlist[index])

                query = 'UPDATE PRODUCTS SET product_code=?, product_name=?, price=?, Category=? WHERE product_code=?'
                values = [code, name, price, arabic_trans(cat), code]
                self.cur.execute(query, values)
                self.conn.commit()

        prodz = self.get_products()
        stocktable = DataTable(table=prodz)
        content.add_widget(stocktable)
        self.update_product_fields()

    def remove_product(self, code):
        content = self.ids.scrn_product_contents
        content.clear_widgets()

        if code == '':
            self.notify.add_widget(Label(text='[color=#FF0000][b]Product code is Required![/b][/color]', markup=True))
            self.notify.open()
            Clock.schedule_once(self.killswitch, 1)
        else:
            products = dict(self.get_products().items())
            cod = products.get('product_code')
            codelist = list(str(a) for a in dict(cod).values())
            x = codelist.__contains__(code)
            if not x:
                self.notify.add_widget(Label(text='[color=#FF0000][b]Invalid Product Code![/b][/color]', markup=True))
                self.notify.open()
                Clock.schedule_once(self.killswitch, 1)
            else:
                query = 'DELETE FROM PRODUCTS WHERE product_code=?'
                values = [code]
                self.cur.execute(query, values)
                self.conn.commit()

        prodz = self.get_products()
        stocktable = DataTable(table=prodz)
        content.add_widget(stocktable)
        self.remove_product_fields()

    def get_products(self):
        _stocks = OrderedDict()
        _stocks['product_code'] = {}
        _stocks['product_name'] = {}
        _stocks['price'] = {}
        _stocks['Category'] = {}

        product_code = []
        product_name = []
        price = []
        cat = []

        query = 'SELECT * FROM PRODUCTS'
        self.cur.execute(query)
        products = self.cur.fetchall()
        for product in products:
            product_code.append(product[1])
            name = product[2]
            '''
            if len(name) > 10:
                name = name[:10] + '...'
            '''
            product_name.append(name)
            price.append(product[3])
            cat.append(product[4])
        products_length = len(product_code)
        idx = 0
        while idx < products_length:
            _stocks['product_code'][idx] = product_code[idx]
            _stocks['product_name'][idx] = product_name[idx]
            _stocks['price'][idx] = price[idx]
            _stocks['Category'][idx] = cat[idx]

            idx += 1

        return _stocks

    # ----------------------------------------------------------------------------------
    # ----------------------------------------------------------------------------------
    # ----------------------------------------------------------------------------------
    # ----------------------------------------------------------------------------------

    def get_customers(self):
        _custs = OrderedDict()
        _custs['customer_code'] = {}
        _custs['customer_name'] = {}
        _custs['phone'] = {}
        _custs['address'] = {}
        _custs['delivery'] = {}

        customer_codes = []
        customer_names = []
        phones = []
        addresses = []
        deliverys = []

        query = 'SELECT * FROM CUSTOMERS'
        self.cur.execute(query)
        customers = self.cur.fetchall()

        for cust in customers:
            customer_codes.append(cust[1])
            customer_names.append(cust[2])
            phones.append(cust[3])
            addresses.append(cust[4])
            deliverys.append(cust[5])
        custs_length = len(customer_codes)
        idx = 0
        while idx < custs_length:
            _custs['customer_code'][idx] = customer_codes[idx]
            _custs['customer_name'][idx] = customer_names[idx]
            _custs['phone'][idx] = phones[idx]
            _custs['address'][idx] = addresses[idx]
            _custs['delivery'][idx] = deliverys[idx]

            idx += 1

        return _custs

    def add_customer_fields(self):
        target = self.ids.ops_fields_s
        target.clear_widgets()
        crud_code = TextInput(hint_text='Customer Code', multiline=False)
        crud_name = Ar_text(hint_text='Customer Name', multiline=False, font_name="utils/STC")
        crud_phone = TextInput(hint_text='Phone', multiline=False, input_filter="int")
        crud_address = Ar_text(hint_text='Address', multiline=False, font_name="utils/STC")
        crud_delivery = TextInput(hint_text='Delivery', multiline=False, input_filter="float")
        crud_submit = Button(text='Add', size_hint_x=None, width=100,
                             on_release=lambda x: self.add_cust(crud_code.text, crud_name.text, crud_phone.text,
                                                                crud_address.text, crud_delivery.text))

        target.add_widget(crud_code)
        target.add_widget(crud_name)
        target.add_widget(crud_phone)
        target.add_widget(crud_address)
        target.add_widget(crud_delivery)
        target.add_widget(crud_submit)

    def update_customer_fields(self):
        target = self.ids.ops_fields_s
        target.clear_widgets()

        crud_code = TextInput(hint_text='Customer Code', multiline=False)
        crud_name = Ar_text(hint_text='Customer Name', multiline=False, font_name="utils/STC")
        crud_phone = TextInput(hint_text='Phone', multiline=False, input_filter="int")
        crud_address = Ar_text(hint_text='Address', multiline=False, font_name="utils/STC")
        crud_delivery = TextInput(hint_text='Delivery', multiline=False, input_filter="float")
        crud_submit = Button(text='Update', size_hint_x=None, width=100,
                             on_release=lambda x: self.update_cust(crud_code.text, crud_name.text, crud_phone.text,
                                                                   crud_address.text, crud_delivery.text))

        target.add_widget(crud_code)
        target.add_widget(crud_name)
        target.add_widget(crud_phone)
        target.add_widget(crud_address)
        target.add_widget(crud_delivery)
        target.add_widget(crud_submit)

    def remove_customer_fields(self):
        target = self.ids.ops_fields_s
        target.clear_widgets()
        crud_code = TextInput(hint_text='Customer Code', multiline=False)
        crud_submit = Button(text='Remove', size_hint_x=None, width=100,
                             on_release=lambda x: self.remove_cust(crud_code.text))

        target.add_widget(crud_code)
        target.add_widget(crud_submit)

    def add_cust(self, code, name, phone, add, delivery):
        content = self.ids.scrn_customer_contents
        content.clear_widgets()

        if code == '' or name == '' or phone == '' or add == '' or delivery == '':
            self.notify.add_widget(Label(text='[color=#FF0000][b]All Fields Required![/b][/color]', markup=True))
            self.notify.open()
            Clock.schedule_once(self.killswitch, 1)
        else:
            query = 'INSERT INTO CUSTOMERS(customer_code, customer_name, phone, address, delivery) VALUES(?, ?, ?, ?, ?)'
            values = [code, arabic_trans(name), phone, arabic_trans(add), delivery]
            self.cur.execute(query, values)
            self.conn.commit()

        custs = self.get_customers()
        custstable = DataTable(table=custs)
        content.add_widget(custstable)
        self.add_customer_fields()

    def update_cust(self, code, name, phone, add, delivery):
        content = self.ids.scrn_customer_contents
        content.clear_widgets()

        if code == '':
            self.notify.add_widget(Label(text='[color=#FF0000][b]Code is Required![/b][/color]', markup=True))
            self.notify.open()
            Clock.schedule_once(self.killswitch, 1)
        else:
            custs = dict(self.get_customers().items())
            codes = custs.get('customer_code')
            codelist = list(dict(codes).values())
            x = codelist.__contains__(code)
            if not x:
                self.notify.add_widget(Label(text='[color=#FF0000][b]Invalid Customer Code![/b][/color]', markup=True))
                self.notify.open()
                Clock.schedule_once(self.killswitch, 1)
            else:
                index = codelist.index(code)
                if name == '':
                    names = custs.get('customer_name')
                    nameslist = list(dict(names).values())
                    name = arabic_trans(nameslist[index])
                if add == '':
                    adds = custs.get('address')
                    addlist = list(dict(adds).values())
                    add = arabic_trans(addlist[index])
                if phone == '':
                    phones = custs.get('phone')
                    phonelist = list(dict(phones).values())
                    phone = phonelist[index]
                if delivery == '':
                    very = custs.get('delivery')
                    verylist = list(dict(very).values())
                    delivery = verylist[index]

                query = 'UPDATE CUSTOMERS SET customer_name=?, phone=?, address=?, delivery=? WHERE customer_code=?'
                values = [arabic_trans(name), phone, arabic_trans(add), delivery, code]
                self.cur.execute(query, values)
                self.conn.commit()

        custs = self.get_customers()
        custstable = DataTable(table=custs)
        content.add_widget(custstable)
        self.update_customer_fields()

    def remove_cust(self, code):
        content = self.ids.scrn_customer_contents
        content.clear_widgets()

        if code == '':
            self.notify.add_widget(Label(text='[color=#FF0000][b]Code is Required![/b][/color]', markup=True))
            self.notify.open()
            Clock.schedule_once(self.killswitch, 1)
        else:
            custs = dict(self.get_customers().items())
            codes = custs.get('customer_code')
            codelist = list(dict(codes).values())
            x = codelist.__contains__(code)
            if not x:
                self.notify.add_widget(Label(text='[color=#FF0000][b]Invalid Customer Code![/b][/color]', markup=True))
                self.notify.open()
                Clock.schedule_once(self.killswitch, 1)
            else:
                query = 'DELETE FROM CUSTOMERS WHERE customer_code=?'
                values = [code]
                self.cur.execute(query, values)
                self.conn.commit()

        custs = self.get_customers()
        custstable = DataTable(table=custs)
        content.add_widget(custstable)
        self.remove_customer_fields()

    # ----------------------------------------------------------------------------------
    # ----------------------------------------------------------------------------------
    # ----------------------------------------------------------------------------------
    # ----------------------------------------------------------------------------------
    def get_orders(self):
        _orders = OrderedDict()
        _orders['ID'] = {}
        _orders['customer_code'] = {}
        _orders['Date'] = {}
        _orders['total_price'] = {}
        _orders['disc'] = {}
        _orders['notes'] = {}

        ids = []
        customer_codes = []
        dates = []
        total_prices = []
        discs = []
        notes = []

        query = 'SELECT * FROM TRANSACTIONS'
        self.cur.execute(query)
        orders = self.cur.fetchall()

        for order in orders:
            ids.append(order[0])
            customer_codes.append(order[1])
            dates.append(order[2])
            total_prices.append(order[3])
            discs.append(order[4])
            notes.append(order[5])
        order_length = len(ids)
        idx = 0
        while idx < order_length:
            _orders['ID'][idx] = ids[idx]
            _orders['customer_code'][idx] = customer_codes[idx]
            _orders['Date'][idx] = dates[idx]
            _orders['total_price'][idx] = total_prices[idx]
            _orders['disc'][idx] = discs[idx]
            _orders['notes'][idx] = notes[idx]

            idx += 1

        return _orders

    def get_order_details(self, id):
        _order_d = OrderedDict()
        _order_d['id'] = {}
        _order_d['order id'] = {}
        _order_d['product name'] = {}
        _order_d['price'] = {}
        _order_d['quantity'] = {}
        _order_d['discount'] = {}
        _order_d['notes'] = {}

        ids = []
        order_ids = []
        prod_ids = []
        prod_names = []
        pnames = list(self.get_products()['product_name'].values())
        prices = []
        quantitys = []
        discounts = []
        notes = []
        query = 'SELECT * FROM TRANS_DETAILS where order_id=' + str(id)
        self.cur.execute(query)
        details = self.cur.fetchall()

        for order in details:
            ids.append(order[0])
            order_ids.append(order[1])
            prod_ids.append(order[2])
            prod_names.append(pnames[list(prod_ids).index(order[2])])
            prices.append(order[3])
            quantitys.append(order[4])
            discounts.append(order[5])
            notes.append(order[6])
        order_length = len(ids)
        idx = 0
        while idx < order_length:
            _order_d['id'][idx] = ids[idx]
            _order_d['order id'][idx] = order_ids[idx]
            _order_d['product name'][idx] = prod_names[idx]
            _order_d['price'][idx] = prices[idx]
            _order_d['quantity'][idx] = quantitys[idx]
            _order_d['discount'][idx] = discounts[idx]
            _order_d['notes'][idx] = notes[idx]

            idx += 1
        print(_order_d)
        return _order_d

    def show_details_fields(self):
        target = self.ids.ops_fields_o
        target.clear_widgets()
        crud_order_id = TextInput(hint_text='order ID', multiline=False, input_filter="int",
                                  on_text_validate=lambda x: self.show_details(crud_order_id.text))
        crud_submit = Button(text='Show', size_hint_x=None, width=100,
                             on_release=lambda x: self.show_details(crud_order_id.text))

        target.add_widget(crud_order_id)
        target.add_widget(crud_submit)

    def update_order_fields(self):
        target = self.ids.ops_fields_o
        target.clear_widgets()
        crud_id = TextInput(hint_text='Order ID', multiline=False, input_filter="int")
        crud_code = TextInput(hint_text='Customer Code', multiline=False)
        crud_total = TextInput(hint_text='Total Price', multiline=False, input_filter="float")
        crud_notes = Ar_text(hint_text='Notes', multiline=False, font_name="utils/STC")
        crud_submit = Button(text='Update', size_hint_x=None, width=100,
                             on_release=lambda x: self.update_order(crud_id.text, crud_code.text, crud_total.text, crud_notes.text,
                                                                    crud_notes.text))

        target.add_widget(crud_id)
        target.add_widget(crud_code)
        target.add_widget(crud_total)
        target.add_widget(crud_notes)
        target.add_widget(crud_submit)

    def remove_order_fields(self):
        target = self.ids.ops_fields_o
        target.clear_widgets()
        crud_id = TextInput(hint_text='order ID', multiline=False, input_filter="int")
        crud_submit = Button(text='Delete', size_hint_x=None, width=100,
                             on_release=lambda x: self.delete_order(crud_id.text))

        target.add_widget(crud_id)
        target.add_widget(crud_submit)

    def show_details(self, order_id):
        if order_id == '':
            self.notify.add_widget(Label(text='[color=#FF0000][b]order ID is Required![/b][/color]', markup=True))
            self.notify.open()
            Clock.schedule_once(self.killswitch, 1)
            return
        details = self.get_order_details(order_id)
        if details['id'].__len__() <= 0:
            self.notify.add_widget(Label(text='[color=#FF0000][b]Invalid order ID[/b][/color]', markup=True))
            self.notify.open()
            Clock.schedule_once(self.killswitch, 1)
            return
        else:
            self.current_order_id_details = order_id
            detailstable = DataTable(table=details)
            self.ids.scrn_mngr.current = 'scrn_order_details'
            self.ids.scrn_order_d_contents.clear_widgets()
            self.ids.scrn_order_d_contents.add_widget(detailstable)
        self.show_details_fields()

    def update_order(self, order_id, cust_code, total, notes):
        content = self.ids.scrn_orders_contents
        content.clear_widgets()

        if order_id == '':
            self.notify.add_widget(Label(text='[color=#FF0000][b]Order ID is Required![/b][/color]', markup=True))
            self.notify.open()
            Clock.schedule_once(self.killswitch, 1)
        else:
            orders = dict(self.get_orders().items())
            ids = orders.get('ID')
            x = False
            if ids is not None:
                order_id = int(order_id)
                idlist = list(dict(ids).values())
                x = idlist.__contains__(order_id)
            if not x:
                self.notify.add_widget(Label(text='[color=#FF0000][b]Invalid Order ID![/b][/color]', markup=True))
                self.notify.open()
                Clock.schedule_once(self.killswitch, 1)
            else:
                index = idlist.index(order_id)
                if cust_code == '':
                    codes = dict(self.get_orders().items()).get('customer_code')
                    codelist = list(dict(codes).values())
                    cust_code = codelist[index]
                if total == '':
                    totals = orders.get('total_price')
                    totallist = list(dict(totals).values())
                    total = totallist[index]
                if notes == '':
                    noteso = orders.get('notes')
                    notelist = list(dict(noteso).values())
                    notes = notelist[index]
                else:
                    notes = arabic_trans(notes)
                query = 'UPDATE TRANSACTIONS SET customer_code=?, total_price=?, notes=? WHERE ID=?'
                values = [cust_code, total, notes, order_id]
                self.cur.execute(query, values)
                self.conn.commit()

        orders = self.get_orders()
        orderstable = DataTable(table=orders)
        content.add_widget(orderstable)
        self.update_order_fields()

    def delete_order(self, order_id):
        content = self.ids.scrn_orders_contents
        content.clear_widgets()

        if order_id == '':
            self.notify.add_widget(Label(text='[color=#FF0000][b]Order ID is Required![/b][/color]', markup=True))
            self.notify.open()
            Clock.schedule_once(self.killswitch, 1)
        else:
            orders = dict(self.get_orders().items())
            ids = orders.get('ID')
            x = False
            if ids is not None:
                codelist = list(dict(ids).values())
                x = codelist.__contains__(int(order_id))
            if not x:
                self.notify.add_widget(Label(text='[color=#FF0000][b]Invalid Order ID![/b][/color]', markup=True))
                self.notify.open()
                Clock.schedule_once(self.killswitch, 1)
            else:
                query = 'DELETE FROM TRANSACTIONS WHERE ID=?'
                values = [order_id]
                self.cur.execute(query, values)
                self.conn.commit()

        orders = self.get_orders()
        orderstable = DataTable(table=orders)
        content.add_widget(orderstable)
        self.remove_order_fields()

    def update_details_fields(self):
        target = self.ids.ops_fields_d
        target.clear_widgets()

        crud_id = TextInput(hint_text='Detail ID', multiline=False, input_filter="int")
        crud_pcode = TextInput(hint_text='Product Code', multiline=False)
        crud_price = TextInput(hint_text='Price', multiline=False, input_filter="float")
        crud_qty = TextInput(hint_text='Quantity', multiline=False, input_filter="float")
        crud_disc = TextInput(hint_text='Discount', multiline=False, input_filter="float")
        crud_note = Ar_text(hint_text=arabic_trans('ملاحظات الصنف'), multiline=False, font_name="utils/STC")
        crud_submit = Button(text='Update', size_hint_x=None, width=100,
                             on_release=lambda x: self.update_details(crud_id.text, crud_pcode.text, crud_price.text,
                                                                      crud_qty.text, crud_disc.text, arabic_trans(crud_note.text)))

        target.add_widget(crud_id)
        target.add_widget(crud_pcode)
        target.add_widget(crud_price)
        target.add_widget(crud_qty)
        target.add_widget(crud_disc)
        target.add_widget(crud_note)
        target.add_widget(crud_submit)

    def add_details_fields(self):
        target = self.ids.ops_fields_d
        target.clear_widgets()

        crud_pcode = TextInput(hint_text='Product Code', multiline=False)
        crud_price = TextInput(hint_text='Price', multiline=False, input_filter="float")
        crud_qty = TextInput(hint_text='Quantity', multiline=False, input_filter="float")
        crud_disc = TextInput(hint_text='Discount', multiline=False, input_filter="float")
        crud_note = Ar_text(hint_text='Notes', multiline=False, input_filter="float", font_name='utils/STC')
        crud_submit = Button(text='Add', size_hint_x=None, width=100,
                             on_release=lambda x: self.add_details(crud_pcode.text, crud_price.text,
                                                                      crud_qty.text, crud_disc.text, crud_note.text))

        target.add_widget(crud_pcode)
        target.add_widget(crud_price)
        target.add_widget(crud_qty)
        target.add_widget(crud_disc)
        target.add_widget(crud_note)
        target.add_widget(crud_submit)

    def add_details(self, pcode, price, qty, disc, note):
        content = self.ids.scrn_order_d_contents
        content.clear_widgets()
        if pcode == '' or price == '' or qty == '' or disc == '' or note == '':
            self.notify.add_widget(Label(text='[color=#FF0000][b]missing data![/b][/color]', markup=True))
            self.notify.open()
            Clock.schedule_once(self.killswitch, 1)
        else:
            qty = float(qty)
            disc = float(disc)
            note = arabic_trans(note)
            query = 'INSERT INTO TRANS_DETAILS(order_id, prod_id, price, quantity, discount, note, date) VALUES(?, ?, ?, ?, ?, ?, ?)'
            values = [self.current_order_id_details, pcode, price, qty, disc, note, datetime.now().date()]
            self.cur.execute(query, values)
            self.conn.commit()

        details = self.get_order_details(self.current_order_id_details)
        detailstable = DataTable(table=details)
        content.add_widget(detailstable)
        self.update_details_fields()

    def update_details(self, id, pcode, price, qty, disc, note):
        content = self.ids.scrn_order_d_contents
        content.clear_widgets()
        if id == '':
            self.notify.add_widget(Label(text='[color=#FF0000][b]Detail ID is Required![/b][/color]', markup=True))
            self.notify.open()
            Clock.schedule_once(self.killswitch, 1)
        else:
            details = dict(self.get_details(id).items())
            ids = details.get('id')
            orderids = details.get('order_id')
            x = False
            if ids is not None:
                id = int(id)
                idlist = list(dict(ids).values())
                orderlist = list(dict(orderids).values())
                index = idlist.index(id)
                x = idlist[index] == int(id) and orderlist[index] == int(self.current_order_id_details)
            if not x:
                self.notify.add_widget(Label(text='[color=#FF0000][b]Invalid detail ID![/b][/color]', markup=True))
                self.notify.open()
                Clock.schedule_once(self.killswitch, 1)
            else:
                if pcode == '':
                    codes = details.get('prod_id')
                    codelist = list(dict(codes).values())
                    pcode = codelist[index]
                if price == '':
                    prices = details.get('price')
                    plist = list(dict(prices).values())
                    price = plist[index]

                if qty == '':
                    qtys = details.get('price')
                    qlist = list(dict(qtys).values())
                    qty = qlist[index]

                if disc == '':
                    discs = details.get('discount')
                    dlist = list(dict(discs).values())
                    disc = dlist[index]

                if note == '':
                    notes = details.get('notes')
                    nlist = list(dict(notes).values())
                    note = nlist[index]

                query = 'UPDATE TRANS_DETAILS SET prod_id=?, price=?, quantity=?, discount=?, note=? WHERE id=?'
                values = [pcode, price, qty, disc, note, id]
                self.cur.execute(query, values)
                self.conn.commit()

        details = self.get_order_details(self.current_order_id_details)
        detailstable = DataTable(table=details)
        content.add_widget(detailstable)
        self.update_details_fields()

    def get_details(self, id):
        _order_d = OrderedDict()
        _order_d['id'] = {}
        _order_d['order_id'] = {}
        _order_d['prod_id'] = {}
        _order_d['price'] = {}
        _order_d['quantity'] = {}
        _order_d['discount'] = {}
        _order_d['notes'] = {}

        ids = []
        order_ids = []
        prod_ids = []
        prices = []
        quantitys = []
        discounts = []
        notes = []
        query = 'SELECT * FROM TRANS_DETAILS where id=' + str(id)
        self.cur.execute(query)
        details = self.cur.fetchall()

        for order in details:
            ids.append(order[0])
            order_ids.append(order[1])
            prod_ids.append(order[2])
            prices.append(order[3])
            quantitys.append(order[4])
            discounts.append(order[5])
            notes.append(order[6])
        order_length = len(ids)
        idx = 0
        while idx < order_length:
            _order_d['id'][idx] = ids[idx]
            _order_d['order_id'][idx] = order_ids[idx]
            _order_d['prod_id'][idx] = prod_ids[idx]
            _order_d['price'][idx] = prices[idx]
            _order_d['quantity'][idx] = quantitys[idx]
            _order_d['discount'][idx] = discounts[idx]
            _order_d['notes'][idx] = notes[idx]

            idx += 1
        return _order_d

    def get_all_details(self):
        _order_d = OrderedDict()
        _order_d['id'] = {}
        _order_d['order_id'] = {}
        _order_d['prod_id'] = {}
        _order_d['price'] = {}
        _order_d['quantity'] = {}
        _order_d['discount'] = {}
        _order_d['notes'] = {}
        _order_d['date'] = {}

        ids = []
        order_ids = []
        prod_ids = []
        prices = []
        quantitys = []
        discounts = []
        dnote = []
        date = []
        query = 'SELECT * FROM TRANS_DETAILS'
        self.cur.execute(query)
        details = self.cur.fetchall()

        for order in details:
            ids.append(order[0])
            order_ids.append(order[1])
            prod_ids.append(order[2])
            prices.append(order[3])
            quantitys.append(order[4])
            discounts.append(order[5])
            dnote.append(order[6])
            date.append(order[7])
        order_length = len(ids)
        idx = 0
        while idx < order_length:
            _order_d['id'][idx] = ids[idx]
            _order_d['order_id'][idx] = order_ids[idx]
            _order_d['prod_id'][idx] = prod_ids[idx]
            _order_d['price'][idx] = prices[idx]
            _order_d['quantity'][idx] = quantitys[idx]
            _order_d['discount'][idx] = discounts[idx]
            _order_d['notes'][idx] = dnote[idx]
            _order_d['date'][idx] = date[idx]

            idx += 1
        return _order_d

    def remove_details_fields(self):
        target = self.ids.ops_fields_d
        target.clear_widgets()
        crud_id = TextInput(hint_text='Detail ID', multiline=False, input_filter="int")
        crud_submit = Button(text='Remove', size_hint_x=None, width=100,
                             on_release=lambda x: self.delete_details(crud_id.text))

        target.add_widget(crud_id)
        target.add_widget(crud_submit)

    def delete_details(self, id):
        content = self.ids.scrn_order_d_contents
        content.clear_widgets()
        if id == '':
            self.notify.add_widget(Label(text='[color=#FF0000][b]Detail ID is Required![/b][/color]', markup=True))
            self.notify.open()
            Clock.schedule_once(self.killswitch, 1)
        else:
            details = dict(self.get_details(id).items())
            ids = details.get('id')
            orderids = details.get('order_id')
            x = False
            if ids is not None:
                id = int(id)
                idlist = list(dict(ids).values())
                orderlist = list(dict(orderids).values())
                index = idlist.index(id)
                x = idlist[index] == int(id) and orderlist[index] == int(self.current_order_id_details)
            if not x:
                self.notify.add_widget(Label(text='[color=#FF0000][b]Invalid detail ID![/b][/color]', markup=True))
                self.notify.open()
                Clock.schedule_once(self.killswitch, 1)
            else:
                query = 'DELETE FROM TRANS_DETAILS WHERE id=' + str(id)
                self.cur.execute(query)
                self.conn.commit()

        details = self.get_order_details(self.current_order_id_details)
        detailstable = DataTable(table=details)
        content.add_widget(detailstable)
        self.remove_details_fields()

    def show_all_orders(self):
        content = self.ids.scrn_orders_contents
        content.clear_widgets()
        orders = self.get_orders()
        orderstable = DataTable(table=orders)
        self.ids.scrn_mngr.current = 'scrn_orders'
        self.ids.scrn_orders_contents.add_widget(orderstable)
        self.ids.ops_fields_o.clear_widgets()
    # ----------------------------------------------------------------------------------
    # ----------------------------------------------------------------------------------
    # ----------------------------------------------------------------------------------
    # ----------------------------------------------------------------------------------

    def get_all_prod_stats_fields(self):
        target = self.ids.ops_fields_stats
        target.clear_widgets()
        date = datetime.today().date().__str__().split('-')
        daysin = TextInput(text=date[2], hint_text='day', multiline=False, input_filter="int")
        monthsin = TextInput(text=date[1], hint_text='month', multiline=False, input_filter="int")
        yearsin = TextInput(text=date[0], hint_text='year', multiline=False, input_filter="int")
        submit = Button(text='Show', size_hint_x=None, width=100, on_release=
        lambda x: self.get_all_prod_stats(yearsin.text, monthsin.text, daysin.text))

        target.add_widget(yearsin)
        target.add_widget(monthsin)
        target.add_widget(daysin)
        target.add_widget(submit)

    def prod_split_fields(self):
        target = self.ids.ops_fields_stats
        target.clear_widgets()
        date = datetime.today().date().__str__().split('-')
        daysin = TextInput(text=date[2], hint_text='day', multiline=False, input_filter="int")
        monthsin = TextInput(text=date[1], hint_text='month', multiline=False, input_filter="int")
        yearsin = TextInput(text=date[0], hint_text='year', multiline=False, input_filter="int")
        prodin = TextInput(hint_text='product code', multiline=False)
        submit = Button(text='Show', size_hint_x=None, width=100, on_release= lambda x: self.prod_split(daysin.text,
                                                                            monthsin.text, yearsin.text, prodin.text))

        target.add_widget(yearsin)
        target.add_widget(monthsin)
        target.add_widget(daysin)
        target.add_widget(prodin)
        target.add_widget(submit)

    def get_all_prod_stats(self, year, month, day):
        if year == '' or month == '' or day == '':
            self.notify.add_widget(
                Label(text='[color=#FF0000][b]Please Enter Missing Information![/b][/color]', markup=True))
            self.notify.open()
            Clock.schedule_once(self.killswitch, 1)
            return

        if len(month) == 1:
            month = '0' + month
        if len(day) == 1:
            day = '0' + day
        date = year + '-' + month + '-' + day
        print(self.dates)
        print(date)
        if date not in self.dates:
            self.notify.add_widget(Label(text='[color=#FF0000][b]Date is Invalid![/b][/color]', markup=True))
            self.notify.open()
            Clock.schedule_once(self.killswitch, 1)
        else:
            target = self.ids.scrn_stats_contents
            target.clear_widgets()
            products = self.get_products()
            prod_names = list(products['product_name'].values())
            prod_ids = list(products['product_code'].values())

            _stats1 = OrderedDict()
            _stats1['product code'] = {}
            _stats1['product name'] = {}
            _stats1['quantity'] = {}

            pcodes = []
            names = []
            quantitys = []
            query = 'SELECT * FROM TRANS_DETAILS where date=\"' + date + '\"'
            self.cur.execute(query)
            prodssplit = self.cur.fetchall()
            for prod in prodssplit:
                pcodes.append(prod[2])
                quantitys.append(prod[4])
                names.append(prod_names[list(prod_ids).index(prod[2])])

            statepcode = []
            statename = []
            statequantity = []
            for i in range(len(pcodes)):
                if pcodes[i] not in statepcode:
                    statepcode.append(pcodes[i])
                    statename.append(names[i])
                    statequantity.append(quantitys[i])
                else:
                    index = statepcode.index(pcodes[i])
                    statequantity[index] += quantitys[i]

            prod_length = len(statepcode)
            idx = 0
            while idx < prod_length:
                _stats1['product code'][idx] = statepcode[idx]
                _stats1['product name'][idx] = statename[idx]
                _stats1['quantity'][idx] = statequantity[idx]
                idx += 1

            statestable = DataTable(table=_stats1)
            target.add_widget(statestable)

    def prod_split(self, day, month, year, prod_id):
        if year == '' or month == '' or day == '' or prod_id == '':
            self.notify.add_widget(
                Label(text='[color=#FF0000][b]Please Enter Missing Information![/b][/color]', markup=True))
            self.notify.open()
            Clock.schedule_once(self.killswitch, 1)
            return
        if len(month) == 1:
            month = '0' + month
        if len(day) == 1:
            day = '0' + day
        date = year + '-' + month + '-' + day
        if date not in self.dates:
            self.notify.add_widget(Label(text='[color=#FF0000][b]Date is Invalid![/b][/color]', markup=True))
            self.notify.open()
            Clock.schedule_once(self.killswitch, 1)
        else:
            target = self.ids.scrn_stats_contents
            target.clear_widgets()
            products = self.get_products()
            prod_names = list(products['product_name'].values())
            prod_ids = list(products['product_code'].values())

            _stats2 = OrderedDict()
            _stats2['product code'] = {}
            _stats2['product name'] = {}
            _stats2['quantity'] = {}
            _stats2['order id'] = {}
            _stats2['notes'] = {}

            name = prod_names[list(prod_ids).index(int(prod_id))]
            quantitys = []
            orderids = []
            notes = []

            query = 'SELECT * FROM TRANS_DETAILS where date=\"' + date + '\" and prod_id=' + str(prod_id)
            self.cur.execute(query)
            prodssplit = self.cur.fetchall()

            for prod in prodssplit:
                quantitys.append(prod[4])
                orderids.append(prod[1])
                notes.append(prod[6])

            prod_length = len(orderids)
            idx = 0
            while idx < prod_length:
                _stats2['product code'][idx] = prod_id
                _stats2['product name'][idx] = name
                _stats2['order id'][idx] = orderids[idx]
                _stats2['quantity'][idx] = quantitys[idx]
                _stats2['notes'][idx] = notes[idx]
                idx += 1

            statestable = DataTable(table=_stats2)
            target.add_widget(statestable)

        # ----------------------------------------------------------------------------------
        # ----------------------------------------------------------------------------------
        # ----------------------------------------------------------------------------------
        # ----------------------------------------------------------------------------------

    def change_screen(self, instance):
        if instance.text == 'Manage Products':
            self.ids.scrn_mngr.current = 'scrn_product_content'

        elif instance.text == 'Manage Users':
            self.ids.scrn_mngr.current = 'scrn_content'

        elif instance.text == 'Manage Customers':
            self.ids.scrn_mngr.current = 'scrn_customer_content'

        elif instance.text == 'Manage Orders':
            self.ids.scrn_mngr.current = 'scrn_orders'
            content = self.ids.scrn_orders_contents
            content.clear_widgets()
            orders = self.get_orders()
            orderstable = DataTable(table=orders)
            content.add_widget(orderstable)

        else:
            self.details = self.get_all_details()
            ids = list(self.details['id'].values())
            self.dates = list()
            self.order_prod_ids = list()
            self.order_prod_names = list()
            for x in range(len(ids)):
                if self.details['date'][x] not in self.dates:
                    self.dates.append(self.details['date'][x])
                    print(self.dates)
                if self.details['prod_id'][x] not in self.order_prod_ids:
                    self.order_prod_ids.append(str(self.details['prod_id'][x]))

                    try:
                        self.order_prod_names.append(self.prodnames[self.prod_ids.index(self.details['prod_id'][x])])
                    except():
                        pass

            product_code = []
            product_name = []
            query = 'SELECT * FROM PRODUCTS'
            self.cur.execute(query)
            products = self.cur.fetchall()
            for product in products:
                product_code.append(product[1])
                name = product[2]
                product_name.append(name)
            self.ids.ops_fields_stats.clear_widgets()
            self.ids.scrn_mngr.current = 'scrn_analysis'

    def print_fields(self):
        target = self.ids.ops_fields_stats
        target.clear_widgets()
        date = datetime.today().date().__str__().split('-')
        daysin = TextInput(text=date[2], hint_text='day', multiline=False, input_filter="int", on_text_validate=lambda x: self.print_reciept(yearsin.text, monthsin.text, daysin.text))
        monthsin = TextInput(text=date[1], hint_text='month', multiline=False, input_filter="int")
        yearsin = TextInput(text=date[0], hint_text='year', multiline=False, input_filter="int")
        submit = Button(text='Print', size_hint_x=None, width=100, on_release=
        lambda x: self.print_reciept(yearsin.text, monthsin.text, daysin.text))

        target.add_widget(yearsin)
        target.add_widget(monthsin)
        target.add_widget(daysin)
        target.add_widget(submit)

    def print_reciept(self, year, month, day):
        if year == '' or month == '' or day == '':
            self.notify.add_widget(
                Label(text='[color=#FF0000][b]Please Enter Missing Information![/b][/color]', markup=True))
            self.notify.open()
            Clock.schedule_once(self.killswitch, 1)
            return
        x = True
        if len(month) == 1:
            month = '0' + month
        if len(day) == 1:
            day = '0' + day
        date = year + '-' + month + '-' + day
        if date not in self.dates:
            self.notify.add_widget(Label(text='[color=#FF0000][b]Date is Invalid![/b][/color]', markup=True))
            self.notify.open()
            Clock.schedule_once(self.killswitch, 1)
        else:
            query = 'SELECT * FROM TRANS_DETAILS WHERE date=\"' + date + '\"'
            self.cur.execute(query)
            details = self.cur.fetchall()

            query = 'SELECT ID FROM TRANSACTIONS WHERE Date=\"' + date + '\"'
            self.cur.execute(query)
            ids = self.cur.fetchall()
            order_ids = []
            for i in ids:
                order_ids.append(i[0])

            products = self.get_products()
            prod_names = list(products['product_name'].values())
            prod_ids = list(products['product_code'].values())
            prices = list(products['price'].values())
            if not os.path.isdir('orders'):
                os.mkdir('orders')
            if not os.path.isdir('orders/' + date):
                os.mkdir('orders/' + date)
            for order in order_ids:
                total = 0.00
                discount = 0.00
                document = Document()
                font = document.styles['Normal'].font
                font.name = 'STC Regular'
                if os.path.exists('logo.jpg'):
                    document.add_picture('logo.jpg', width=Inches(1))
                p = document.add_paragraph('')
                p.add_run('Order ID: ' + str(order)).bold = True
                table = document.add_table(rows=1, cols=5)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'الاجمالي'
                hdr_cells[1].text = 'الكمية'
                hdr_cells[2].text = 'السعر'
                hdr_cells[4].text = 'الصنف'
                hdr_cells[3].text = 'ملحوظه'
                for detail in details:
                    detail = list(detail)

                    if order == detail[1]:
                        pname = prod_names[prod_ids.index(detail[2])]
                        uprice = prices[prod_ids.index(detail[2])]
                        row_cells = table.add_row().cells
                        row_cells[4].text = arabic_trans(arabic_trans(arabic_trans(pname)))
                        row_cells[1].text = str(detail[4])
                        row_cells[2].text = str(uprice)
                        row_cells[0].text = str(detail[3]) + 'ج'
                        if detail[6] != '' and detail[6] is not None:
                            row_cells[3].text = arabic_trans(detail[6])
                        total += detail[3]
                        discount += detail[5]
                for row in table.rows:
                    row.height = Cm(0.1)
                query = 'SELECT customer_code, notes  FROM TRANSACTIONS WHERE ID=' + str(order)
                self.cur.execute(query)
                [cust_code, notes] = self.cur.fetchone()

                document.add_paragraph(
                    '__________________________________________').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                if cust_code != 'new':
                    query = 'SELECT * FROM CUSTOMERS WHERE customer_code=\"' + str(cust_code) + '\"'
                    self.cur.execute(query)
                    cust = self.cur.fetchone()
                else:
                    cust = (0, 'new', 'new', 'new', 'new', 5)
                if discount > 0.00:
                    document.add_paragraph(str(discount) + 'ج ' + '\t\t\t:' + 'خصم ').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                document.add_paragraph(str(cust[5]) + 'ج ' + '\t\t\t\t\t\t\t\t\t:' + 'خدمة التوصيل ').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                document.add_paragraph(str(total + cust[5]) + 'ج ' + '\t\t\t\t\t\t\t\t\t\t:' +  'اجمالي ').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if cust[3] == 'new':
                    document.add_paragraph(':' + 'العنوان').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    document.add_paragraph(':' + 'رقم الهاتف: ').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    document.add_paragraph(arabic_trans(cust[4]) + '\t\t\t:' + 'العنوان ').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    document.add_paragraph(cust[3] + '\t\t\t:' + 'رقم الهاتف ').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if notes:
                    document.add_paragraph(arabic_trans(notes) + '\t\t\t:' + 'الملحوظات ').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                try:
                    document.save('orders/' + date + '/order-' + str(order) + '.docx')

                except:
                    x = False

            if not x:
                self.notify.add_widget(
                    Label(text='[color=#FF0000][b]Close the open order document to save![/b][/color]', markup=True))
                self.notify.open()
                Clock.schedule_once(self.killswitch, 10)
            else:
                self.notify.add_widget(
                    Label(text='[color=#00FF00][b]Print Orders is Complete![/b][/color]', markup=True))
                self.notify.open()
                Clock.schedule_once(self.killswitch, 10)
# ----------------------------------------------------------------------------------
# ----------------------------------------------------------------------------------
# ----------------------------------------------------------------------------------
# ----------------------------------------------------------------------------------


class AdminApp(App):
    def build(self):
        return AdminWindow()


if __name__ == '__main__':
    AdminApp().run()